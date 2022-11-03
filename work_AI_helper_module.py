# ...\ Вначале был закомментированный участок кода

# + + + библиотеки
import re
import requests as rq
from bs4 import BeautifulSoup as bs
import docx
import time
import subprocess
import os
import configparser


# + + + функции
def kolvo_zn_p_zpt(num, zn=0):
    return f"{num:.{zn}f}"  # превращает число в строку, быть осторожным лучше заменить на округление
    # а это оставить для печатной формы, чтобы .0 превращать в .00


def proverka_vvoda(chek_prav_option, text):  # проверяет правильно ли ввел ответ пользователь
    print("Результат поиска верный? (Введите + или - в ответ)")
    while True:
        rezult = input("Ваш выбор: ").strip()
        rezult = unputFilter(rezult, "?")
        if rezult == "+":
            print("Замечательно, продолжим.")
            chek_prav_option = True
            break
        elif rezult == "-":
            print(f"Хм.. тогда прошу вас проверить {text}, и я попробую поискать еще раз.")
            break
        else:
            print("А что вы ввели? Попробуйте еще раз... ")
            continue
    return chek_prav_option


def str_to_num_transform(x):  # занимается преобразованием строк в флоат с коррекцией точек и запятых при вводе пользователем разных типов

    num_transform = x
    num_transform_list = num_transform.split()

    y = len(num_transform_list)
    i = 1

    while i != y:
        num_transform_list[0] = num_transform_list[0] + num_transform_list[i]
        i += 1
    if "," in num_transform_list[0]:
        x = float(num_transform_list[0].replace(',', '.'))
    else:
        x = float(num_transform_list[0])
    return x


def list_iterator(lst, n):  # эта функция берет список подготовленный для деления по заданному количеству элементов для того, чтобы создать
    return [lst[i:i + n] for i in range(0, len(lst), n)]  # новый список где каждый элемент это спискок из этого количества элементов первоначального списка


def is_in_lst(str_, words):
    for word in words:
        if word.lower() in str_.lower():
            return True
    return False


def unputFilter(text, option):
    if option == "nu":
        alphabet = '0123456789'
    elif option == "pr":
        alphabet = '0123456789.,'
    elif option == "?":
        alphabet = '+-*'
    rezultat = [c for c in text if c in alphabet]
    return ''.join(rezultat)


def createConfig(path):  # Замут с ини - создание

    config = configparser.ConfigParser()
    config.add_section("Settings")
    config.set("Settings", "antidemp_koef", "0.25")
    config.set("Settings", "okopfus", "X")
    config.set("Settings", "podpisant_us", "X")
    config.set("Settings", "delayer", "1")
    config.set("Settings", "fio_gen_dir_us", "X")
    config.set("Settings", "nds_pro", "20")
    config.set("Settings", "antidemp_koef_mnoj", "1.5")
    config.set("Settings", "shtraf_smp_pro", "1")
    config.set("Settings", "shtraf_nesmp_pro", "10")
    config.set("Settings", "flag_nesmp_ot_cena_maksimal", "Нет")
    config.set("Settings", "flag_smp_ot_cena_gk", "Да")
    config.set("Settings", "flag_smp_strf_b5km1k", "Да")
    config.set("Settings", "site_for_parsing", "http:")
    config.set("Settings", "site_for_parsing_short", ":")
    config.set("Settings", "igk_gk4", "...")
    config.set("Settings", "site_for_parsing_firm", "http:")
    config.set("Settings", "rekviz_org_us", "Пока не заполнено")
    config.set("Settings", "dop_url1", "http:")
    config.set("Settings", "dop_url2", "http:")
    config.set("Settings", "dop_url3", "http:")
    config.set("Settings", "dop_url_comm", "...")
    config.set("Settings", "dop_url4", "http:")

    with open(path, "w") as config_file:
        config.write(config_file)


def readConfig1(path):  # Замут с ини - чтение

    config = configparser.ConfigParser()
    config.read(path)

    antidemp_koef = config.get("Settings", "antidemp_koef")
    okopfus = config.get("Settings", "okopfus")
    gen_dir_us = config.get("Settings", "podpisant_us")
    delayer = config.get("Settings", "delayer")
    fio_gen_dir_us = config.get("Settings", "fio_gen_dir_us")

    return antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us


def readConfig2(path):  # Замут с ини - чтение 2

    config = configparser.ConfigParser()
    config.read(path)

    antidemp_koef_mnoj = config.get("Settings", "antidemp_koef_mnoj")
    shtraf_smp_pro = config.get("Settings", "shtraf_smp_pro")
    shtraf_nesmp_pro = config.get("Settings", "shtraf_nesmp_pro")

    return antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro


def readConfig3(path):  # Замут с ини - чтение 3

    config = configparser.ConfigParser()
    config.read(path)

    nds_pro = config.get("Settings", "nds_pro")
    flag_nesmp_ot_cena_maksimal = config.get("Settings", "flag_nesmp_ot_cena_maksimal")
    flag_smp_ot_cena_gk = config.get("Settings", "flag_smp_ot_cena_gk")
    flag_smp_strf_b5km1k = config.get("Settings", "flag_smp_strf_b5km1k")

    return nds_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k


def readConfig_fst_run(path):

    config = configparser.ConfigParser()
    config.read(path)

    site_for_parsing = config.get("Settings", "site_for_parsing")
    site_for_parsing = str(site_for_parsing)

    site_for_parsing_short = config.get("Settings", "site_for_parsing_short")
    site_for_parsing_short = str(site_for_parsing_short)

    igk_gk4 = config.get("Settings", "igk_gk4")
    igk_gk4 = str(igk_gk4)

    site_for_parsing_firm = config.get("Settings", "site_for_parsing_firm")
    site_for_parsing_firm = str(site_for_parsing_firm)

    rekviz_org_us = config.get("Settings", "rekviz_org_us")
    rekviz_org_us = str(rekviz_org_us)

    return site_for_parsing, site_for_parsing_short, igk_gk4, site_for_parsing_firm, rekviz_org_us


def readConfig_fst_run_dop_urls(path):
  
    config = configparser.ConfigParser()
    config.read(path)

    dop_url1 = config.get("Settings", "dop_url1")
    dop_url1 = str(dop_url1)

    dop_url2 = config.get("Settings", "dop_url2")
    dop_url2 = str(dop_url2)

    dop_url3 = config.get("Settings", "dop_url3")
    dop_url3 = str(dop_url3)

    dop_url_comm = config.get("Settings", "dop_url_comm")
    dop_url_comm = str(dop_url_comm)

    dop_url4 = config.get("Settings", "dop_url4")
    dop_url4 = str(dop_url4)

    return dop_url1, dop_url2, dop_url3, dop_url_comm, dop_url4


def writeConfig(path, nomer, znachenie):  # Замут с ини - изменение 2
    config = configparser.ConfigParser()
    config.read(path)

    if nomer == "1":
        config.set("Settings", "antidemp_koef", znachenie)
    elif nomer == "2":
        config.set("Settings", "delayer", znachenie)
    elif nomer == "3":
        config.set("Settings", "okopfus", znachenie)
    elif nomer == "4":
        config.set("Settings", "podpisant_us", znachenie)
    elif nomer == "5":
        config.set("Settings", "fio_gen_dir_us", znachenie)
    elif nomer == "6":
        config.set("Settings", "nds_pro", znachenie)
    elif nomer == "7":
        config.set("Settings", "antidemp_koef_mnoj", znachenie)
    elif nomer == "8":
        config.set("Settings", "shtraf_smp_pro", znachenie)
    elif nomer == "9":
        config.set("Settings", "shtraf_nesmp_pro", znachenie)
    elif nomer == "10":
        config.set("Settings", "flag_nesmp_ot_cena_maksimal", znachenie)
    elif nomer == "11":
        config.set("Settings", "flag_smp_ot_cena_gk", znachenie)
    elif nomer == "12":
        config.set("Settings", "flag_smp_strf_b5km1k", znachenie)
    elif nomer == "13":
        config.set("Settings", "site_for_parsing", znachenie)
    elif nomer == "14":
        config.set("Settings", "site_for_parsing_short", znachenie)
    elif nomer == "15":
        config.set("Settings", "igk_gk4", znachenie)
    elif nomer == "16":
        config.set("Settings", "site_for_parsing_firm", znachenie)
    elif nomer == "17":
        config.set("Settings", "rekviz_org_us", znachenie)
    elif nomer == "18":
        config.set("Settings", "dop_url1", znachenie)
    elif nomer == "19":
        config.set("Settings", "dop_url2", znachenie)
    elif nomer == "20":
        config.set("Settings", "dop_url3", znachenie)
    elif nomer == "21":
        config.set("Settings", "dop_url_comm", znachenie)
    elif nomer == "22":
        config.set("Settings", "dop_url4", znachenie)
    else:
        print("Элемент отсутствует в списке.")

    with open(path, "w") as config_file:
        config.write(config_file)


def poisk_nomenklaturi_v_tablicax(x):
    return x.has_attr("rowspan")


def transform_perem_types(antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k):
    antidemp_koef = float(antidemp_koef)
    okopfus = str(okopfus)
    gen_dir_us = str(gen_dir_us)
    delayer = int(delayer)
    fio_gen_dir_us = str(fio_gen_dir_us)
    nds_pro = float(nds_pro)
    antidemp_koef_mnoj = float(antidemp_koef_mnoj)
    shtraf_smp_pro = float(shtraf_smp_pro)
    shtraf_nesmp_pro = float(shtraf_nesmp_pro)
    flag_nesmp_ot_cena_maksimal = str(flag_nesmp_ot_cena_maksimal)
    flag_smp_ot_cena_gk = str(flag_smp_ot_cena_gk)
    flag_smp_strf_b5km1k = str(flag_smp_strf_b5km1k)
    return antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k


# + + + переменные
internet = False; rejim_vibran = False; rejim_online = False; reprovinternet = "*"; flag_gk_bez_obiema = False
chek_prav_option1 = False; chek_prav_option2 = False; gos_oboron_zakaz = False; smp_check = False; igk_gk = None
vopros_s_viborom_ck = False; vopros_s_viborom_ck_so_zv = False; nds_vbr_flag = False
specifikaciya = []
telef_org = None; ogrnip_org = None; kpp_org = None; ogrn_org = None; ur_adres_org = None; pcht_adres_org = None
okpo = None; oktmo = None; okopf = None; rasch_schet = None; korr_schet = None; imya_banka = None; bik = None
slvr_zagr_url = {}
athef4dwnld = []
new_str = " "
nomer_dla_opredelenia_slovarnix_ssilok = 1
critikal_Error = False
exitProgramm = False
path_conf = "settings.ini"


# + + + параметры
headers = {
    "Accept": "*/*",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.84 Safari/537.36 OPR/85.0.4341.72"
}

# + + + + + + + + + + + Начало программы + + + + + + + + + + +
intro = """+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++gen+1+++version+1.21+β++++ +
+                                                                                                           +
+                                   Module for Automatic Intelligence -                                     +
+                                           помощник по заполнению контрактов                               +
+                                                                                                           +
+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++by+SQR++++++ +"""
# + + +
print(intro)
print(new_str)
print("Приветствую! ")
time.sleep(1)
print(new_str)
print("Я помогу вам заполнить контракт. \nПожалуйста, следуйте инструкциям, которые я буду выводить на экран.")

if not os.path.exists(path_conf):
    print(new_str)
    print("Мной обнаружено отсутствие конфигурационного файла... ")
    time.sleep(1)
    print("Сейчас создам его с настройками по умолчанию.")
    createConfig(path_conf)
    time.sleep(2)
    print("Готово.")
    print(new_str)
    print("Для корректной работы программы необходимо дополнить конфигурационный файл.")
    time.sleep(1)
    print("Это также необходимо для настройки алгоритма под пользователя.")
    time.sleep(1)
    print(new_str)
    print("Первое, с чего мы начнем - это основной сайт для поиска данных. Укажите его ниже: ")
    site_for_parsing = str(input(">>> "))
    writeConfig(path_conf, "13", site_for_parsing)
    site_for_parsing_short_lst = list(site_for_parsing)
    del(site_for_parsing_short_lst[0:8])
    site_for_parsing_short = "".join(site_for_parsing_short_lst)
    writeConfig(path_conf, "14", site_for_parsing_short)
    time.sleep(1)
    print(new_str)
    print("Далее укажите окончание ИГК контрактов вашей организации (14-25 разряды): ")
    igk_gk4 = str(input(">>> "))
    writeConfig(path_conf, "15", igk_gk4)
    time.sleep(1)
    print(new_str)
    print("Хорошо. Теперь укажите сайт для поиска данных по организации: ")
    site_for_parsing_firm = str(input(">>> "))
    writeConfig(path_conf, "16", site_for_parsing_firm)
    time.sleep(1)
    print(new_str)
    print("Так, а сейчас введите несколько персонализирующих значений: ")
    time.sleep(1)
    print("1. Текст по форме: в лице *должность* *директор*, действующего на основании *документ/доверенность + дата*")
    podpisant_us = str(input(">>> "))
    writeConfig(path_conf, "4", podpisant_us)
    time.sleep(1)
    print(new_str)
    print("2. Инициалы директора - Фамилия И.О.: ")
    fio_gen_dir_us = str(input(">>> "))
    writeConfig(path_conf, "5", fio_gen_dir_us)
    time.sleep(1)
    print(new_str)
    print("3. ОКОПФ вашей организации: ")
    okopfus = str(input(">>> "))
    writeConfig(path_conf, "3", okopfus)
    time.sleep(1)
    print(new_str)
    print("На этом настройка конфигурационного файла завершена! ")
    time.sleep(1)
    print("Дополнительно эти и другие данные можно изменить или дополнить в разделе настройки. ")
    print(new_str)
    time.sleep(1)
    print("Запускаю основной алгоритм...")
    time.sleep(2)
    print("Готово!")
else:
    time.sleep(1)
    print(new_str)


# получаем переменные из конфига
antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us = readConfig1(path_conf)
antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro = readConfig2(path_conf)
nds_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k = readConfig3(path_conf)

# трансформируем значения в нужные типы данных, а также присваиваем переменным значения из модуля выше
antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, \
shtraf_smp_pro, shtraf_nesmp_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, \
flag_smp_strf_b5km1k = transform_perem_types(antidemp_koef, okopfus, gen_dir_us, delayer, \
fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro, \
flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k)

site_for_parsing, site_for_parsing_short, igk_gk4, site_for_parsing_firm, rekviz_org_us = readConfig_fst_run(path_conf)
dop_url1, dop_url2, dop_url3, dop_url_comm, dop_url4 = readConfig_fst_run_dop_urls(path_conf)


print("Для продолжения просто нажмите клавишу [ENTER] ")
print(new_str)
print("А если вы хотите вызвать справку или меню настроек, то:")
print(" -> Для вызова справки (инструкции по программе) введите слово \' Инструкция \' и нажмите клавишу [ENTER]")
print("   -> Для редактирования настроек введите слово \' Настройка \' и нажмите клавишу [ENTER]")


while rejim_vibran != True:
    if rejim_online != False:
        print(new_str)
        print("Также повторюсь: если вы хотите сразу начать просто нажмите клавишу [ENTER]")
        print("            |.... Для вызова справки (инструкции по программе) введите слово \' Инструкция \' и нажмите клавишу [ENTER]")
        print("                |.... Для редактирования настроек введите слово \' Настройка \' и нажмите клавишу [ENTER]")
        rejim_online = False
    print(new_str)
    startoption = input("Ожидаю ввод тут:").lower().strip()
    print(new_str)
    if startoption == "инструкция" or startoption == "справка":
        print("""+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ +
+                                                                                                           +
+                                                  СПРАВКА                                                  +

          1. Данный модуль (приложение) написан с целью упрощения и автоматизации рутины по заполнению
          Контрактов
          2. Когда приложение ожидает что-либо от пользователя, она обычно дает об этом знать с помощью   
          инструкции на экране. Ожидая ввод, она в зависимости от заданного вопроса ждет, что пользователь
          введет знак + или - или * или какое-то текстовое\цифровое значение, а потом нажмет клавишу
          [ENTER]
          3. Текст в окне приложения можно выделять и копировать стандартным сочетанием клавиш CTRL + C     
          4. Текст в поле для ввода можно вставлять стандартным сочетанием клавиш CTRL + V  
          5. Как правило в функционал приложения встроены первичные анализаторы ввода пользователя, и
          при неправильном вводе будет запрошен повторный ввод правильного значения. Однако в определенных
          случаях может понадобиться закрыть программу и начать заполнение заново.
          6. В разделе Настройки есть возможность менять некоторые значения по умолчанию, необходимые для
          работы приложения. Их перечень можно узнать на экране настроек. Их следует поменять если, 
          к примеру, изменилось законодательство или что-либо еще, что делает работу приложения
          неактуальной. 
          7. Приложение может создавать файл с расширением INI. Его следует хранить в той же папке что и
          EXE файл. Это файл с сохраненными настройками.
          8. Приложение не будет работать без подключения к интернету.
          9. Приложение создаст ворд файл с названием > Собранная информация по закупке.docx < 
          в этом файле будут отражены все результаты работы приложения
          10. Некоторую информацию приложение не сможет подгрузить в результате работы. Но оставит 
          ценные ссылочки, где ее можно взять.  
          11. Приложение скачает также и все файлы извещения, такие как проект контракта, ТЗ, расчет 
          НМЦК и любые другие, если сочтет их нужными.
          12. Приложение не вмешивается в текст файлов по ряду технических причин, а также по причине 
          возможного НЕПРЕДНАМЕРЕННОГО ИСКАЖЕНИЯ текста контракта, что может повлечь юридические
          последствия.       
          13. Для эффективной работы следует внимательно читать и выполнять инструкции представляемые 
          приложением. В случае вылетов, странного поведения приложения, ошибок и других не типичных
          явлений - стоит попробовать осуществить запуск приложения еще раз и попытаться верно 
          выполнять инструкции. В случае повторения ошибок, следует перейти к заполнению контракта по 
          необходимой закупке вручную. А если полетели настройки, можно вручную удалить файл
          settings.ini - приложение создаст его с настройками по умолчанию.                                                       
+                                                                                                           +
+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ +
        """)
        print("Повторюсь: если вы хотите сразу начать просто нажмите клавишу [ENTER]")
        print("        |.... Для вызова справки (инструкции по программе) введите слово \' Инструкция \' и нажмите клавишу [ENTER]")
        print("            |.... Для редактирования настроек введите слово \' Настройка \' и нажмите клавишу [ENTER]")
    elif startoption == "настройка":
        print("+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ +")
        print("                                               Опции")
        print("+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ +")
        print(f"""
        Отладочная информация:

            1.  Порог включения антидемпинговых мер: {antidemp_koef}%
            2.  Время задержки вывода: {delayer} сек 
            3.  ОКОПФ: {okopfus}
            4.  Заказчик: {gen_dir_us} 
            5.  Подписант: {fio_gen_dir_us} 
            6.  Процент НДС: {nds_pro} %
            7.  Множитель антидемпинговых мер: {antidemp_koef_mnoj}
            8.  Размер штрафа для СМП: {shtraf_smp_pro} %
            9.  Размер штрафа для не СМП: {shtraf_nesmp_pro} %
            10. Штраф у не СМП считается от НМЦК: {flag_nesmp_ot_cena_maksimal}
            11. Штраф у СМП считается от Цены контракта: {flag_smp_ot_cena_gk}
            12. Правило расчета штрафов у СМП (> 1000 но < 5000) актуально: {flag_smp_strf_b5km1k}  
            13. Основной сайт для поиска информации по закупке: {site_for_parsing}
            14. Его сокращенная версия: {site_for_parsing_short}
            15. Окончание ИГК для организации : {igk_gk4}
            16. Сайт для поиска информации о контрагенте: {site_for_parsing_firm}
            17. Реквизиты организации: {rekviz_org_us}
            18. Доп. сайт по организациям 1: {dop_url1}
            19. Доп. сайт по организациям 2: {dop_url2}
            20. Доп. сайт по организациям 3: {dop_url3}
            21. Указание на сайт площадки: {dop_url_comm}
            22. Доп. сайт по организациям 4: {dop_url4}
        """)
        print("+ +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ +")
        print(" ")
        stopOptions = False
        while stopOptions == False:
            print("При необходимости отредактировать опцию введите ее номер и нажмите [ENTER], для выхода из настроек просто нажмите [ENTER]")
            print(new_str)
            optionsProg = input("Ввод номера опции: ").strip()
            if optionsProg == "1":
                print(f"Порог включения антидемпинговых мер: {antidemp_koef}")
                optionsProg1 = input("Введите нужное значение (любое положительное дробное число ЧЕРЕЗ ТОЧКУ, а не запятую): ").strip()
                if "," in optionsProg1:
                    optionsProg1 = optionsProg1.replace(",", ".")
                writeConfig(path_conf, "1", optionsProg1)
            elif optionsProg == "2":
                print(f"Время задержки вывода сообщений: {delayer}")
                optionsProg2 = input("Введите нужное значение (целое число от 0 до разумных пределов): ").strip()
                writeConfig(path_conf, "2", optionsProg2)
            elif optionsProg == "3":
                print(f"ОКОПФ: {okopfus}")
                optionsProg3 = input("Введите нужное значение: ").strip()
                writeConfig(path_conf, "3", optionsProg3)
            elif optionsProg == "4":
                print(f"{gen_dir_us}")
                optionsProg4 = input("Введите нужное значение (Скопируйте полностью фразу и измените в ней что-либо при неободимости): ").strip()
                writeConfig(path_conf, "4", optionsProg4)
            elif optionsProg == "5":
                print(f"{fio_gen_dir_us}")
                optionsProg5 = input("Введите нужное значение (Скопируйте полностью текст и измените его при неободимости): ").strip()
                writeConfig(path_conf, "5", optionsProg5)
            elif optionsProg == "6":
                print(f"Процент НДС: {nds_pro}")
                optionsProg6 = input("Введите нужное значение (Число, без знака > % <, дробную часть (при наличии) отделите точкой, а не запятой): ").strip()
                if "," in optionsProg6:
                    optionsProg6 = optionsProg6.replace(",", ".")
                writeConfig(path_conf, "6", optionsProg6)
            elif optionsProg == "7":
                print(f"Множитель антидемпинговых мер: {antidemp_koef_mnoj}")
                optionsProg7 = input("Введите нужное значение (Число, дробную часть (при наличии) отделите точкой, а не запятой): ").strip()
                if "," in optionsProg7:
                    optionsProg7 = optionsProg7.replace(",", ".")
                writeConfig(path_conf, "7", optionsProg7)
            elif optionsProg == "8":
                print(f"Размер штрафа для СМП: {shtraf_smp_pro}")
                optionsProg8 = input("Введите нужное значение (Число, без знака > % <, дробную часть (при наличии) отделите точкой, а не запятой): ").strip()
                if "," in optionsProg8:
                    optionsProg8 = optionsProg8.replace(",", ".")
                writeConfig(path_conf, "8", optionsProg8)
            elif optionsProg == "9":
                print(f"Размер штрафа для не СМП: {shtraf_nesmp_pro}")
                optionsProg9 = input("Введите нужное значение (Число, без знака > % <, дробную часть (при наличии) отделите точкой, а не запятой): ").strip()
                if "," in optionsProg9:
                    optionsProg9 = optionsProg9.replace(",", ".")
                writeConfig(path_conf, "9", optionsProg9)
            elif optionsProg == "10":
                print(f"Штраф у не СМП считается от НМЦК: {flag_nesmp_ot_cena_maksimal}")
                optionsProg10 = input("Введите нужное значение (Да или Нет): ").strip()
                writeConfig(path_conf, "10", optionsProg10)
            elif optionsProg == "11":
                print(f"Штраф у СМП считается от Цены контракта: {flag_smp_ot_cena_gk}")
                optionsProg11 = input("Введите нужное значение (Да или Нет): ").strip()
                writeConfig(path_conf, "11", optionsProg11)
            elif optionsProg == "12":
                print(f"Правило расчета штрафов у СМП (> 1000 но < 5000) актуально: {flag_smp_strf_b5km1k}")
                optionsProg12 = input("Введите нужное значение (Да или Нет): ").strip()
                writeConfig(path_conf, "12", optionsProg12)
            elif optionsProg == "13":
                print(f"Основной сайт для поиска информации по закупке: {site_for_parsing}")
                optionsProg13 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "13", optionsProg13)
            elif optionsProg == "14":
                print(f"Сокращенная версия основного сайта из п.13: {site_for_parsing_short}")
                optionsProg14 = input("Введите нужное значение (без http и www): ").strip()
                writeConfig(path_conf, "14", optionsProg14)
            elif optionsProg == "15":
                print(f"Окончание ИГК контрактов вашей организации (14-25 разряды): {igk_gk4}")
                optionsProg15 = input("Введите нужное значение (число, в соотв. с требованиями организации): ").strip()
                writeConfig(path_conf, "15", optionsProg15)
            elif optionsProg == "16":
                print(f"Сайт для поиска информации о контрагенте: {site_for_parsing_firm}")
                optionsProg16 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "16", optionsProg16)
            elif optionsProg == "17":
                print(f"Реквизиты организации, которые используются в контрактах: {rekviz_org_us}")
                optionsProg17 = input("Введите нужное значение (текстом, без окопф): ").strip()
                writeConfig(path_conf, "17", optionsProg17)
            elif optionsProg == "18":
                print(f"Доп. сайт по организациям 1: {dop_url1}")
                optionsProg18 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "18", optionsProg18)
            elif optionsProg == "19":
                print(f"Доп. сайт по организациям 2: {dop_url2}")
                optionsProg19 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "19", optionsProg19)
            elif optionsProg == "20":
                print(f"Доп. сайт по организациям 3: {dop_url3}")
                optionsProg20 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "20", optionsProg20)
            elif optionsProg == "21":
                print(f"Указание на сайт площадки: {dop_url_comm}")
                optionsProg21 = input("Введите нужное значение (текстом): ").strip()
                writeConfig(path_conf, "21", optionsProg21)
            elif optionsProg == "22":
                print(f"Доп. сайт по организациям 4: {dop_url4}")
                optionsProg22 = input("Введите нужное значение (http...../): ").strip()
                writeConfig(path_conf, "22", optionsProg22)
            else:
                stopOptions = True
            print(new_str)


        antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us = readConfig1(path_conf)
        antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro = readConfig2(path_conf)
        nds_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k = readConfig3(path_conf)

        antidemp_koef, okopfus, gen_dir_us, delayer, fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, \
        shtraf_smp_pro, shtraf_nesmp_pro, flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, \
        flag_smp_strf_b5km1k = transform_perem_types(antidemp_koef, okopfus, gen_dir_us, delayer, \
        fio_gen_dir_us, nds_pro, antidemp_koef_mnoj, shtraf_smp_pro, shtraf_nesmp_pro, \
        flag_nesmp_ot_cena_maksimal, flag_smp_ot_cena_gk, flag_smp_strf_b5km1k)

        site_for_parsing, site_for_parsing_short, igk_gk4, site_for_parsing_firm, rekviz_org_us = readConfig_fst_run(path_conf)
        dop_url1, dop_url2, dop_url3, dop_url_comm, dop_url4 = readConfig_fst_run_dop_urls(path_conf)


        print(new_str)
        print("Повторюсь: если вы хотите сразу начать просто нажмите клавишу [ENTER]")
        print("        |.... Для вызова справки (инструкции по программе) введите слово \' Инструкция \' и нажмите клавишу [ENTER]")
        print("            |.... Для редактирования настроек введите слово \' Настройка \' и нажмите клавишу [ENTER]")
    else:
        print("Принято!")
        print(new_str)
        print("Сейчас начнем работу.")
        print("И для этого я проверю интернет соединение, пожалуйста, подождите...")
        print(new_str)

        while rejim_online != True:

            try:
                subprocess.check_call(["ping", "77.88.8.8"], stdout=subprocess.DEVNULL, stderr=subprocess.STDOUT)
                print("Ага... Сетевое подключение работает. \nНаличие доступа в интернет подтверждаю. ")
                internet = True

            except subprocess.CalledProcessError:
                print("Кажется, ваш компьютер не подключен к интернету. \nПроверить еще раз? (где + да, проверить; - нет, не нужно)")
                print(new_str)
                time.sleep(delayer)
                reprovinternet = input("Ожидаю ввод тут: ").strip()
                reprovinternet = unputFilter(reprovinternet, "?")

            if internet != True:
                print(new_str)

                if reprovinternet == "+":
                    print("Проверяю еще раз...")
                elif reprovinternet == "-":
                    print("Тогда возращаемся к началу. Вы также можете закрыть окно, нажав на [X] в верхнем правом углу.")
                    rejim_online = True
                else:
                    print("Вы ввели неверное значение. Предлагаю вернуться в начало или закрыть окно, нажав на [X] в верхнем правом углу.")
                    rejim_online = True

            else:
                rejim_online = True
                rejim_vibran = True

print(new_str)
print("Итак, приступим.")
time.sleep(delayer)
print(new_str)
print("Первое, что мне понадобится, это номер закупки.")
print(new_str)


while chek_prav_option1 != True:
    nomer_zakup = input("Введите его здесь: ").strip()
    nomer_zakup = unputFilter(nomer_zakup, "nu")
    url = f'{site_for_parsing}epz/order/notice/ea20/view/common-info.html?regNumber={nomer_zakup}'
    r = rq.get(url, timeout=10, headers=headers)
    html = bs(r.content, "html.parser")
    object_zakup = html.find(text=re.compile("объект")).next_element.next_element.text.strip()
    print("Найдена закупка: " + object_zakup)
    chek_prav_option1 = proverka_vvoda(chek_prav_option1, "номер закупки")

# Запросы к страничке карточки электронного аукциона после одобрения варианта
cena_maksimal = html.find("span", text=re.compile("аксимальн")).next_element.next_element.next_element.text.strip()
if html.find('span', text=re.compile("евозможно определить количеств")) != None:
    flag_gk_bez_obiema = True
    c_ed_gk1 = html.find('span', text=re.compile("альная сумма цен единиц")).next_element. \
        next_element.next_element.get_text().strip()
    c_ed_gk_lst = c_ed_gk1.split(" ")
    c_ed_gk_lst = list(filter(None, c_ed_gk_lst))
    if len(c_ed_gk_lst) >= 3:
        print(new_str)
        print("Возникла ошибка расчета начальной суммы цен единиц номенклатур. Очищаю ошибочное значение...")
        c_ed_gk1 = 0
        print("Готово. Во избежание неверного подсчета цены Контракта перепроверьте правильность цен вручную! ")
        print(new_str)
    else:
        c_ed_gk1 = c_ed_gk_lst[0]
    if "\xa0" or "\n" in c_ed_gk1:
        if "\xa0" and "\n":
            c_ed_gk = re.sub("[\xa0|\n]", "", c_ed_gk1)
        if "\n" in c_ed_gk1:
            c_ed_gk = re.sub("[\n]", "", c_ed_gk1)
        if "\xa0" in c_ed_gk1:
            c_ed_gk = re.sub("[\xa0]", "", c_ed_gk1)
    else:
        c_ed_gk = c_ed_gk1

ikz = html.find(text=re.compile("ИКЗ")).next_element.next_element.text.strip()
smp_preimuzhj = html.find("span", text=re.compile("ч. 3 ст. 30 Закона"))
obespech_proc = html.find("span", text=re.compile("азмер обеспечения исполнения")).next_element.next_element.next_element.text.strip()
if flag_gk_bez_obiema == True or '₽' in obespech_proc:
    obespech_proc_lst = obespech_proc.split("₽")
    obespech_proc = int(obespech_proc_lst[1].replace("(", "").replace(")", "").replace("%", "").strip())
else:
    obespech_proc = int(obespech_proc.replace(" ", "").replace("%", "").strip())
kbk = html.find("div", id=re.compile("budgetTableInnerHtml")).find("table", class_="blockInfo__table tableBlock").find("td", class_="table__row-item normal-text tableBlock__col_left").get_text().strip()

# После чего переход к страничке результатов торгов и запросы к ней
url_sub = f'{site_for_parsing}epz/order/notice/ea20/view/supplier-results.html?regNumber={nomer_zakup}'
r_sub = rq.get(url_sub, timeout=10, headers=headers)
html_sub = bs(r_sub.content, "html.parser")
data_results = html_sub.find(text=re.compile("Дата")).next_element.next_element.text.strip()
data_results_list = data_results.split()

if html_sub.find(text=re.compile("обедитель")) != None:
    ck_predv = html_sub.find(text=re.compile("обедитель")).next_element.next_element.text.strip()
elif html_sub.find(text=re.compile("срока подачи заявок не подано ни одной заявки")) != None:
    ck_predv = 0
    print("Анализ показал, что закупка не состоялась!")
    critikal_Error = True
else:
    ck_predv = html_sub.find("div", id=re.compile("supplier-def-result-participant-table")) \
        .find("table", class_="blockInfo__table tableBlock").find("tbody").find("td") \
        .next_element.next_element.next_element.next_element.next_element.next_element.text.strip()

while critikal_Error != True or exitProgramm != True:
    if '18803020840692019215' in kbk or '18803020830292019214' in kbk or '18803020830292019211' in kbk:
        gos_oboron_zakaz = True
        print(new_str)
        print("Анализ показал, что закупка проводилась в рамках ГОЗ.")
        time.sleep(delayer)
        print("Поэтому предлагаю сразу сформировать ИГК к данной закупке.")
        print(new_str)
        time.sleep(delayer)
        print("Введите последние две цифры года заключения государственного контракта")
        igk_gk1 = input("Ожидаю значение: ").strip()
        igk_gk1 = unputFilter(igk_gk1, "nu")
        print("Введите последние две цифры года окончания срока действия государственного контракта")
        igk_gk2 = input("Ожидаю значение: ").strip()
        igk_gk2 = unputFilter(igk_gk2, "nu")
        print("""Введите одно число в зависимости от вашего случая:
        1 - Если аукцион состоялся (2 и более участников торговались)
        2 - Если не состоялся (1 участник)
        3 - Если это заключение контракта с ед. поставщиком        
        """)
        igk_gk3 = input("Ожидаю значение: ").strip()
        igk_gk3 = unputFilter(igk_gk3, "nu")
        igk_gk = igk_gk1 + igk_gk2 + "188" + igk_gk3 + igk_gk4

    print(new_str)
    print("""Теперь, пожалуйста, введите ИНН контрагента.
    """)

    while chek_prav_option2 != True and critikal_Error != True:
        inn_org = input("Ожидаю ввод тут: ").strip()
        inn_org = unputFilter(inn_org, "nu")
        url_inn = f'{site_for_parsing}epz/eruz/search/results.html?searchString={inn_org}'
        r_inn = rq.get(url_inn, timeout=10, headers=headers)
        html_inn = bs(r_inn.content, "html.parser")
        data_results_inn = html_inn.find("a", text=re.compile("№")).text.strip()
        data_results_inn = data_results_inn.split(" ")
        url_inn_respond = f'{site_for_parsing}epz/eruz/card/general-information.html?reestrNumber={data_results_inn[1]}'
        r_inn_respond = rq.get(url_inn_respond, timeout=10, headers=headers)
        html_inn_respond = bs(r_inn_respond.content, "html.parser")
        imya_org = html_inn_respond.find(text=re.compile("Участник закупки")).next_element.next_element.get_text().strip().lower().title()
        print("Найден контрагент: " + imya_org)
        chek_prav_option2 = proverka_vvoda(chek_prav_option2, "ИНН участника")

    print(new_str)
    # Запросы к страничке организации после одобрения варианта
    ip_ooo = html_inn_respond.find(text=re.compile("Тип участника закупки")).next_element.next_element.get_text().strip()
    if "ридическ" in ip_ooo:
        UR_lic = True
        sokr_imya_org = html_inn_respond.find(text=re.compile("окращенное наименование")).next_element.next_element.get_text().strip()
        kpp_org = html_inn_respond.find(text=re.compile("КПП")).next_element.next_element.get_text().strip()
        ogrn_org = html_inn_respond.find(text=re.compile("ОГРН")).next_element.next_element.get_text().strip()
        smp_prover = html_inn_respond.find(text=re.compile("частник закупки является субъектом")).next_element.next_element.next_element.get_text().strip()
        ur_adres_org = html_inn_respond.find(text=re.compile("места нахождения")).next_element.next_element.get_text().strip()
        ur_adres_org = ur_adres_org.lower().title()
        fio_gendir = html_inn_respond.find("tbody", class_="tableBlock__body").find("td", class_="tableBlock__col").text
        fio_gendir = fio_gendir.lower().title()
        doljn_gendir = html_inn_respond.find("tbody", class_="tableBlock__body").find("td", class_="tableBlock__col").next_element.next_element.next_element.text
        doljn_gendir = doljn_gendir.lower()
        pcht_adres_org = html_inn_respond.find(text=re.compile("очтовый адрес")).next_element.next_element.get_text().strip()
        pcht_adres_org = pcht_adres_org.lower().title()
        email_org = html_inn_respond.find(text=re.compile("дрес электронной почты")).next_element.next_element.get_text().strip()
        telef_org = html_inn_respond.find(text=re.compile("онтактный телефон")).next_element.next_element.get_text().strip()
    elif "ндивидуальный" in ip_ooo and gos_oboron_zakaz != True:
        UR_lic = False
        fio_gendir = imya_org.lower().title()
        sokr_imya_org = "ИП " + fio_gendir
        imya_org = "Индивидуальный предприниматель " + fio_gendir
        ogrnip_org = html_inn_respond.find(text=re.compile("ОГРНИП")).next_element.next_element.get_text().strip()
        ogrnip_org_data = html_inn_respond.find(text=re.compile("ата регистрации индивидуального предпринимателя")).next_element.next_element.get_text().strip()
        email_org = html_inn_respond.find(text=re.compile("дрес электронной почты")).next_element.next_element.get_text().strip()
        try:
            smp_prover = html_inn_respond.find(text=re.compile("частник закупки является субъектом")).next_element.next_element.next_element.get_text().strip()
        except:
            smp_prover = False
    elif gos_oboron_zakaz == True and "ндивидуальный" in ip_ooo:
        print("Заключение Контракта в рамках ГОЗ невозможно с ИП и физ. лицами. Проверьте правильность данных!")
        critikal_Error = True
        break
    elif gos_oboron_zakaz == True and "изическое лицо" in ip_ooo:
        print("Заключение Контракта в рамках ГОЗ невозможно с ИП и физ. лицами. Проверьте правильность данных!")
        critikal_Error = True
        break
    else:
        gos_oboron_zakaz = True
        print("Искомый контрагент, скорее всего, физическое лицо или присутствует ошибка в алгоритме поиска даннных.")
        time.sleep(delayer)
        print("Пожалуйста закройте программу и попробуйте заново.")

    try:
        url_dop_info_company = f'{site_for_parsing_firm}search?query={inn_org}&search_inactive=0'
        r_dop_info_company = rq.get(url_dop_info_company, timeout=10, headers=headers)
        html_dop_info_company = bs(r_dop_info_company.content, "html.parser")
        okpo = html_dop_info_company.find('span', id="clip_okpo").text.strip()
        oktmo = html_dop_info_company.find('span', id="clip_oktmo").text.strip()
        okopf = html_dop_info_company.find('span', id="clip_okopf").text.strip()
        rasch_schet = 'Не найден'
        korr_schet = 'Не найден'
        imya_banka = 'Не найден'
        bik = 'Не найден'
    except:
        okpo = 'Не найден'
        oktmo = 'Не найден'
        okopf = 'Не найден'
        rasch_schet = 'Не найден'
        korr_schet = 'Не найден'
        imya_banka = 'Не найден'
        bik = 'Не найден'

    # парсинг товарного листа или списка закупаемых работ / услуг
    print("Пытаюсь выгрузить перечень номенклатуры объекта закупки...")
    time.sleep(delayer)
    try:
        url_pechat_form = f'{site_for_parsing}epz/order/notice/printForm/view.html?regNumber={nomer_zakup}'
        r_pechat_form = rq.get(url_pechat_form, timeout=10, headers=headers)
        html_pechat_form = bs(r_pechat_form.content, "html.parser")
        tablitca_cen = html_pechat_form.find_all(poisk_nomenklaturi_v_tablicax)

        for i in tablitca_cen:  # добавить из всех строк и столбцов таблицы данные о товарах услугах или работах и их ценах в список
            if "th" not in str(i) and "border-top" not in str(i):

                info = i.get_text()
                info = info.strip()

                if info != "":
                    specifikaciya.append(info)

        for i in specifikaciya:  # цикл для того чтобы поубирать пробелы из цифр и имен подтянутые из другой кодировки
            if "\xa0" in i:
                specifikaciya[specifikaciya.index(i)] = specifikaciya[specifikaciya.index(i)].replace('\xa0', '')
            if "\n" in i:
                specifikaciya[specifikaciya.index(i)] = specifikaciya[specifikaciya.index(i)].replace('\n', '')

        spisok_cen_objecta_zakup = list_iterator(specifikaciya, 6)  # полученный список преобразован в список из списков
        for i in spisok_cen_objecta_zakup:  # по 6 элементов, которые являют собой данные об одной позиции
            i.pop(1)

        print("Готово. Выгрузка прошла успешно.")
        print(new_str)
        time.sleep(delayer)
    except:
        print("Выгрузка окончилась рядом ошибок. Очищаю ошибочные данные...")
        spisok_cen_objecta_zakup = ["Товары, работы, услуги - не выгружены."]
        print("Готово.")
        print(new_str)
        time.sleep(delayer)

    # Сбор доп инфы - отмена переходим к ручному редактированию в полученном ворде
    print("Итак, основная информация собрана. ")
    time.sleep(delayer)
    print("Далее - позвольте уточнить нюансы. ")

    if flag_gk_bez_obiema == True:
        ck = cena_maksimal
    else:
        print("Начнем с формированию Цены контракта.")
        print(new_str)
        time.sleep(delayer)
        print("Подскажите, будет ли увеличение цены контракта до НМЦК?")
        time.sleep(delayer)
        print("""
        В ответ введите один из следующих символов:
                + (что значит - да, доводим до НМЦК)
                - (что значит - нет, не доводим)
                * (что значит - доводим, но частично - сумма будет между результатом
                                                       аукциона и НМЦК)""")

        while vopros_s_viborom_ck != True:
            ck_vbr = input("Ваш выбор: ").strip()
            ck_vbr = unputFilter(ck_vbr, "?")
            if ck_vbr == "-":
                ck = ck_predv
                print("Принято!")
                vopros_s_viborom_ck = True
            elif ck_vbr == "*":
                print("Тогда введите, пожалуйста, Цену контракта (числом без пробелов, копейки при наличии отделяйте точкой)")
                while vopros_s_viborom_ck_so_zv != True:
                    ck = input("Ожидаю ввод здесь: ").strip()
                    ck = unputFilter(ck, "pr")
                    ck_predv_test = ck_predv
                    cena_maksimal_test = cena_maksimal

                    if " " in ck:
                        ck = ck.replace(" ", "")
                    if " " in ck_predv_test:
                        ck_predv_test = ck_predv_test.replace(" ", "")
                    if " " in cena_maksimal_test:
                        cena_maksimal_test = cena_maksimal_test.replace(" ", "")
                    if "," in ck:
                        ck = ck.replace(",", ".")
                    if "," in ck_predv_test:
                        ck_predv_test = ck_predv_test.replace(",", ".")
                    if "," in cena_maksimal_test:
                        cena_maksimal_test = cena_maksimal_test.replace(",", ".")
                    if "\xa0" in cena_maksimal_test:
                        cena_maksimal_test = cena_maksimal_test.replace("\xa0", "")

                    try:
                        ck_test = float(ck)
                    except:
                        ck_test = 0

                    if ck_test > float(ck_predv_test) and ck_test < float(cena_maksimal_test):
                        print("Принято!")
                        print("Цена ГК устанавливается в размере: " + kolvo_zn_p_zpt(ck_test, 2))
                        vopros_s_viborom_ck = True
                        vopros_s_viborom_ck_so_zv = True
                    else:
                        print("Неправильное число, попробуйте ввести данные еще раз!")
                        print("Проверьте цену, чтобы она была меньше НМЦК, но больше цены предложенной победителем.")

            elif ck_vbr == "+":
                ck = cena_maksimal
                print("Принято!")
                vopros_s_viborom_ck = True
            else:
                print("Введенные вами данные некорректны, попробуйте еще раз!")

    # начало расчета суммы контракта и спецификации
    ck = str_to_num_transform(ck)
    cena_maksimal = str_to_num_transform(cena_maksimal)
    ck_predv = str_to_num_transform(ck_predv)

    print(new_str)
    time.sleep(delayer)
    print("Подскажите, является ли организация плательщиком НДС?")
    time.sleep(delayer)
    url_inn1 = f'{site_for_parsing}epz/contract/search/results.html?morphology=on&fz44=on&contractStageList_0=on&contractStageList_1=on&contractStageList_2=on&contractStageList_3=on&contractStageList=0%2C1%2C2%2C3&selectedContractDataChanges=ANY&contractCurrencyID=-1&budgetLevelsIdNameHidden=%7B%7D&supplierTitle={inn_org}&countryRegIdNameHidden=%7B%7D&sortBy=UPDATE_DATE&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false'
    print("Кстати, это проверить еще можно с помощью поиска вот тут: " + url_inn1)
    print("""
    Но вернемся к вопросу: 
        + значит да, контрагент платит 20% НДС;
        - значит нет, конрагент на упрощенной системе налогообложения
    """)

    while nds_vbr_flag != True:
        nds_vbr = input("Ваш выбор: ").strip()
        nds_vbr = unputFilter(nds_vbr, "?")
        if nds_vbr == "-":
            nds = "Без НДС"
            ysn = True
            nds_vbr_flag = True
        elif nds_vbr == "+":
            nds = ck * nds_pro / (100 + nds_pro)
            nds = round(nds, 2)
            ysn = False
            nds_vbr_flag = True
        else:
            print("Что это значит? Попробуйте еще раз...")
            print(new_str)

    if smp_preimuzhj != None:
        smp_check = True

    try:
        if flag_gk_bez_obiema == True:
            c_ed_gk = str_to_num_transform(c_ed_gk)
            koefauk = ck_predv / c_ed_gk
        else:
            koefauk = ck_predv / cena_maksimal
    except:
        koefauk = None

    try:
        if gos_oboron_zakaz == True or flag_gk_bez_obiema == True or smp_check == False:
            obespech_gk = cena_maksimal * (obespech_proc / 100)
        else:
            obespech_gk = ck * (obespech_proc / 100)
        obespech_gk = round(obespech_gk, 2)
    except:
        obespech_gk = None

    print(new_str)
    print("Угу. Спасибо за сведения! ")
    time.sleep(delayer)
    print("Теперь я проведу обобщение, структурирование и анализ собранной информации...")
    print(new_str)

    # уточняем переменные
    podpisant1 = gen_dir_us

    if UR_lic == True:
        podpisant2 = f"в лице {doljn_gendir} {fio_gendir}, действующего на основании Устава"
    else:
        podpisant2 = f"в лице индивидуального предпринимателя {fio_gendir}, действующего на основании ОГРНИП: {ogrnip_org} от {ogrnip_org_data} года"

    # Создаем файл ворд с найденной инфой
    document = docx.Document()
    document.add_heading('Собранные данные о закупке', 0)
    document.add_paragraph(f'Объект: {object_zakup}')
    document.add_paragraph(f'ИКЗ: {ikz}')
    if gos_oboron_zakaz == True:
        document.add_paragraph(f'ИГК: {igk_gk}, где вместо 00ХХ нужно написать порядковый номер закупки по ГОЗ по журналу')
    document.add_paragraph(f'Подписант 1: {podpisant1}')
    document.add_paragraph(f'Подписант 2: {podpisant2}')
    document.add_paragraph(f'Протокол подведения итогов № {nomer_zakup} от {data_results_list[0]} г.')
    document.add_paragraph(f'КБК для данной закупки будет необходимо отразить в Контракте следующим: {kbk}')
    document.add_paragraph(f'НМЦК: ' + kolvo_zn_p_zpt(cena_maksimal, 2) + ' ₽')
    document.add_paragraph(f'Цена контракта: ' + kolvo_zn_p_zpt(ck, 2) + ' ₽')
    if ysn == True:
        document.add_paragraph(f'Контракт не предусматривает расчет НДС')
    else:
        document.add_paragraph(f'НДС: ' + kolvo_zn_p_zpt(nds, 2) + ' ₽')
    document.add_paragraph(f'Коэффициент ценового падения по аукциону равняется: {koefauk}')

    try:
        if (1 - koefauk) > antidemp_koef and koefauk != 1.0:
            document.add_paragraph(f'Это больше, чем {antidemp_koef}. Поэтому к закупке должны быть применены антидемпинговые меры в части обеспечения.')
            obespech_gk = obespech_gk * antidemp_koef_mnoj
            obespech_gk = round(obespech_gk, 2)
            antidemp = True
        else:
            antidemp = False
            document.add_paragraph(f'К закупке антидемпинговые меры в части обеспечения не применяются.')
        document.add_paragraph('Обеспечение контракта составит: ' + kolvo_zn_p_zpt(obespech_gk, 2) + ' ₽')

        if smp_check == True:
            obespech_gk_fraza = "от цены, по которой заключается контракт"
        else:
            obespech_gk_fraza = "от НМЦК"

        if flag_gk_bez_obiema == True:
            obespech_gk_fraza = "от максимального значения цены контракта"

        if antidemp == True:  # момент про что составляет 30% цены контракта
            document.add_paragraph(f'-> что составляет сумму из расчета {obespech_proc}% {obespech_gk_fraza}, увеличенной в 1,5 раза.')
        else:
            document.add_paragraph(f'-> что составляет сумму из расчета {obespech_proc}% {obespech_gk_fraza}.')
    except:
        obespech_proc = "не выполнен"
        document.add_paragraph(f'Расчет обеспечения ГК {obespech_proc} из-за ошибки в алогритме.')

    try:
        if smp_check == True:

            if flag_smp_ot_cena_gk == "Да":
                shtraf = ck * shtraf_smp_pro / 100
                shtraf = round(shtraf, 2)
                document.add_paragraph(f'Анализ показал, что закупка проводилась среди СМП, поэтому штрафы рассчитывались как {shtraf_smp_pro}% от цены контракта. ')
            elif flag_smp_ot_cena_gk == "Нет":
                shtraf = cena_maksimal * shtraf_smp_pro / 100
                shtraf = round(shtraf, 2)
                document.add_paragraph(f'Анализ показал, что закупка проводилась среди СМП, поэтому штрафы рассчитывались как {shtraf_smp_pro}% от НМЦК. ')

            if flag_smp_strf_b5km1k == "Да":
                if shtraf > 5000.0:
                    shtraf = 5000.0
                elif shtraf < 1000.0:
                    shtraf = 1000.0
        else:

            if flag_nesmp_ot_cena_maksimal == "Да":
                shtraf = cena_maksimal * shtraf_nesmp_pro / 100
                shtraf = round(shtraf, 2)
                document.add_paragraph(f'Анализ показал, что закупка не проводилась среди СМП, поэтому штраф рассчитывался как {shtraf_nesmp_pro}% от НМЦК.')
            elif flag_nesmp_ot_cena_maksimal == "Нет":
                shtraf = ck * shtraf_nesmp_pro / 100
                shtraf = round(shtraf, 2)
                document.add_paragraph(f'Анализ показал, что закупка не проводилась среди СМП, поэтому штраф рассчитывался как {shtraf_nesmp_pro}% от Цены контракта.')

        document.add_paragraph('В связи с чем, штраф для контрагента по условиям контракта составит: ' + kolvo_zn_p_zpt(shtraf, 2) + ' ₽')

    except:
        shtraf = "не вычислен"
        document.add_paragraph(f'Штраф {shtraf} из-за ошибки в алгоритме.')

    document.add_paragraph('Не забудьте проверить в реквизитах вашей организации наличие следующего элемента: ')
    document.add_paragraph(f'>>> ОКОПФ: {okopfus} <<<')
    document.add_paragraph(f' ')
    document.add_paragraph(f' ')
    document.add_paragraph(f' ')

    document.add_paragraph('Реквизиты контрагента:')
    document.add_paragraph(f'Полное наименование: {imya_org}')
    document.add_paragraph(f'Сокращенный вариант: {sokr_imya_org}')
    document.add_paragraph(f'Юридический адрес: {ur_adres_org}')
    document.add_paragraph(f'Почтовый адрес: {pcht_adres_org}')
    document.add_paragraph(f'Контактный телефон: {telef_org}')
    document.add_paragraph(f'e-mail: {email_org}')
    document.add_paragraph(f'ИНН: {inn_org}')
    document.add_paragraph(f'КПП: {kpp_org}')

    if "ридическ" in ip_ooo:
        document.add_paragraph(f'ОГРН: {ogrn_org}')
    if "ндивидуальный" in ip_ooo:
        document.add_paragraph(f'ОГРНИП: {ogrnip_org}')

    table = document.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Государственный заказчик:'
    hdr_cells[1].text = 'Контрагент:'
    hdr_cells1 = table.rows[1].cells

    hdr_cells1[0].add_paragraph(f"""{rekviz_org_us}
ОКОПФ: {okopfus}. """)

    hdr_cells1[1].add_paragraph(f"""{sokr_imya_org}
Юридический адрес: {ur_adres_org}
Почтовый адрес: {pcht_adres_org}
Контактный телефон: {telef_org}
e-mail: {email_org}""")
    hdr_cells1[1].add_paragraph(f"""Получатель:
ИНН: {inn_org}
КПП: {kpp_org}""")

    if "ридическ" in ip_ooo:
        hdr_cells1[1].add_paragraph(f'ОГРН: {ogrn_org}')
    if "ндивидуальный" in ip_ooo:
        hdr_cells1[1].add_paragraph(f'ОГРНИП: {ogrnip_org}')

    hdr_cells1[1].add_paragraph(f"""ОКПО: {okpo}
ОКТМО: {oktmo}
ОКОПФ: {okopf}
р/с: {rasch_schet}
к/с: {korr_schet}
Наименование банка: {imya_banka}
БИК: {bik}""")

    document.add_paragraph(f' ')

    table1 = document.add_table(rows=2, cols=2)
    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = 'Государственный заказчик:'
    hdr1_cells[1].text = 'Контрагент:'
    hdr1_cells1 = table1.rows[1].cells
    hdr1_cells1[0].add_paragraph(f'________________ / {fio_gen_dir_us}/')
    fio_gendir_4red = f"{fio_gendir}"
    fio_gendir_inicial = fio_gendir_4red.split(" ")
    hdr1_cells1[1].add_paragraph(f'________________ / {fio_gendir_inicial[0]} {fio_gendir_inicial[1][0]}.{fio_gendir_inicial[2][0]}./')

    document.add_paragraph(f' ')

    document.add_paragraph(f'Дополнительную информацию и реквизиты контрагента вы можете найти по следующим ссылкам: ')
    document.add_paragraph(f'{dop_url1}search?query={inn_org}', style='List Bullet')
    document.add_paragraph(f'{dop_url2}search?type=inn&val={inn_org}', style='List Bullet')
    document.add_paragraph(f'{dop_url3}search?query={inn_org}&search_inactive=0', style='List Bullet')
    document.add_paragraph(f'На сайте {dop_url_comm} посмотрите в заявке участника его банковские реквизиты:')
    document.add_paragraph(f'{dop_url4}customer/lk/Auctions/Table/', style='List Bullet')

    document.add_paragraph(f' ')
    document.add_paragraph(f' ')

    if spisok_cen_objecta_zakup == ["Товары, работы, услуги - не выгружены."] or spisok_cen_objecta_zakup == []:
        document.add_paragraph('Товары, работы, услуги - не выгружены.')
    else:
        document.add_paragraph('Справочно: расчет цены по перечню товаров, работ, услуг')
        document.add_paragraph('В данной таблице расчет ндс, расчетные показатели по перечню и т.д. - выполнены исходя из недоведения до НМЦК! Используйте их как примерный подсчет')
        kolvo_nomenkl = len(spisok_cen_objecta_zakup)
        table2 = document.add_table(rows=kolvo_nomenkl + 2, cols=8)
        hdr2_cells = table2.rows[0].cells
        hdr2_cells[0].text = 'Номер'
        hdr2_cells[1].text = 'Наименование'
        hdr2_cells[2].text = 'Кол-во / Ед. изм.'
        hdr2_cells[3].text = 'НДС'
        hdr2_cells[4].text = 'Цена за ед.'
        hdr2_cells[5].text = 'Стоимость'
        hdr2_cells[6].text = 'Цена за ед. расчетная (цена ед * коэф аук)'
        hdr2_cells[7].text = 'Стоимость расчетная (цена ед расч * кол-во)'

        nomer_strok_tabl = 1
        for i in spisok_cen_objecta_zakup:
            cells2 = table2.rows[nomer_strok_tabl].cells
            cells2[0].add_paragraph(f'{nomer_strok_tabl}')
            cells2[1].add_paragraph(f'{i[0]}')
            cells2[2].add_paragraph(f'{i[2]} {i[1]}')
            if ysn == True:
                cells2[3].add_paragraph(f'Без НДС')
            else:
                try:
                    nds_nomenclaturi = float(i[3]) * koefauk
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    nds_nomenclaturi = float(nds_nomenclaturi) * nds_pro / (100 + nds_pro)
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    nds_nomenclaturi = nds_nomenclaturi * float(i[2])
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    nds_nomenclaturi = kolvo_zn_p_zpt(nds_nomenclaturi, 2)
                    cells2[3].add_paragraph(f'{nds_nomenclaturi}')
                except:
                    cells2[3].add_paragraph('Ошибка')
            try:
                i3 = kolvo_zn_p_zpt(float(i[3]), 2)
                cells2[4].add_paragraph(f'{i3}')
            except:
                try:
                    cells2[4].add_paragraph(f'{i[3]}')
                except:
                    cells2[4].add_paragraph('Ошибка')
            try:
                i4 = kolvo_zn_p_zpt(float(i[4]), 2)
                cells2[5].add_paragraph(f'{i4}')
            except:
                try:
                    cells2[5].add_paragraph(f'{i[4]}')
                except:
                    cells2[5].add_paragraph('Ошибка')
            try:
                cena_ed_posle_auka = float(i[3]) * koefauk
                cena_ed_posle_auka = round(cena_ed_posle_auka, 2)
                cena_ed_posle_auka = kolvo_zn_p_zpt(cena_ed_posle_auka, 2)
                cells2[6].add_paragraph(f'{cena_ed_posle_auka}')
            except:
                cells2[6].add_paragraph('Ошибка')
            try:
                stoim_nomenkl_posle_auka = float(cena_ed_posle_auka) * float(i[2])
                stoim_nomenkl_posle_auka = round(stoim_nomenkl_posle_auka, 2)
                stoim_nomenkl_posle_auka = kolvo_zn_p_zpt(stoim_nomenkl_posle_auka, 2)
                cells2[7].add_paragraph(f'{stoim_nomenkl_posle_auka}')
            except:
                cells2[7].add_paragraph('Ошибка')
            nomer_strok_tabl += 1

        cellsLST = table2.rows[-1].cells
        cellsLST[0].text = ''
        cellsLST[1].text = 'Итоги'
        summa_po_tablice_k = 0
        for i in spisok_cen_objecta_zakup:
            summa_po_tablice_k += float(i[2])
        summa_po_tablice_k = kolvo_zn_p_zpt(summa_po_tablice_k, 2)
        cellsLST[2].text = f'{summa_po_tablice_k}'
        if ysn == True:
            cellsLST[3].add_paragraph(f'Без НДС')
        else:
            summa_po_tablice_nsd = 0
            try:
                for i in spisok_cen_objecta_zakup:
                    nds_nomenclaturi = float(i[3]) * koefauk
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    nds_nomenclaturi = float(nds_nomenclaturi) * nds_pro / (100 + nds_pro)
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    nds_nomenclaturi = nds_nomenclaturi * float(i[2])
                    nds_nomenclaturi = round(nds_nomenclaturi, 2)
                    summa_po_tablice_nsd += nds_nomenclaturi
                summa_po_tablice_nsd = kolvo_zn_p_zpt(summa_po_tablice_nsd, 2)
                cellsLST[3].add_paragraph(f'{summa_po_tablice_nsd}')
            except:
                cellsLST[3].add_paragraph('Ошибка')
        cellsLST[4].text = ' '
        summa_po_tablice_s1 = 0
        try:
            for i in spisok_cen_objecta_zakup:
                summa_po_tablice_s1 += float(i[4])
                summa_po_tablice_s1 = round(summa_po_tablice_s1, 2)
            summa_po_tablice_s1 = kolvo_zn_p_zpt(summa_po_tablice_s1, 2)
            cellsLST[5].text = f'{summa_po_tablice_s1}'
        except:
            cellsLST[5].text = 'Ошибка'
        cellsLST[6].text = ' '
        summa_po_tablice_s2 = 0
        try:
            for i in spisok_cen_objecta_zakup:
                cena_ed_posle_auka_s2 = float(i[3]) * koefauk
                cena_ed_posle_auka_s2 = round(cena_ed_posle_auka_s2, 2)
                cena_ed_posle_auka_s2 = float(cena_ed_posle_auka_s2) * float(i[2])
                cena_ed_posle_auka_s2 = round(cena_ed_posle_auka_s2, 2)
                summa_po_tablice_s2 += cena_ed_posle_auka_s2
            summa_po_tablice_s2 = kolvo_zn_p_zpt(summa_po_tablice_s2, 2)
            cellsLST[7].text = f'{summa_po_tablice_s2}'
        except:
            cellsLST[7].text = 'Ошибка'

    document.save('Собранная информация по закупке.docx')

    print("Готово! ")
    print(new_str)
    time.sleep(delayer)
    print(f"А я сейчас попробую получить файлы извещения для заполнения.\nПожалуйста, подождите...")
    try:
        zagr_url = f"{site_for_parsing}epz/order/notice/ea20/view/documents.html?regNumber={nomer_zakup}"
        r_zagr = rq.get(zagr_url, timeout=10, headers=headers)
        html_zagr = bs(r_zagr.content, "html.parser")
        athef4dwnldall = html_zagr.find("div", class_=re.compile("first-row-active-documents")) \
            .find_all("a", title=re.compile(""))
        slvr_zagr_url = {}
        athef4dwnld = []
        for i in athef4dwnldall:
            i = f"""{i}"""
            if f"{site_for_parsing_short}44fz/filestore/public/1.0/download/priz/file.html?uid" in i:
                athef4dwnld.append(i.split('"'))
        for i in athef4dwnld:
            slvr_zagr_url[i[3]] = i[1]

        proverka_nalichia = 0
        poiskNEtrKzayavke = ['ребовани', 'аявк']

        for i in slvr_zagr_url:
            i = str(i)

            if is_in_lst(i, poiskNEtrKzayavke) == False:
                dwld_url_gk = f'{slvr_zagr_url[i]}'
                r_dwld_url_gk = rq.get(dwld_url_gk, timeout=10, headers=headers)
                with open(f'{i}', 'wb') as f:
                    f.write(r_dwld_url_gk.content)
                proverka_nalichia += 1

        if proverka_nalichia < 3:
            print("Подгружаю все файлы извещения...")
            for i in slvr_zagr_url:
                dwld_url_gk = f'{slvr_zagr_url[i]}'
                r_dwld_url_gk = rq.get(dwld_url_gk, timeout=10, headers=headers)
                with open(f'{i}', 'wb') as f:
                    f.write(r_dwld_url_gk.content)

        print(new_str)
        print("Выгрузка окончилась успешно! ")

    except:
        print(new_str)
        print("Выгрузка окончилась рядом ошибок.  ")
        print("Пожалуйста, скачайте файлы извещения вручную. ")
        print("Также вы можете обратиться с номером данной закупки к моему разработчику.")
        print("Проверяю доступные задачи для выполнения... ")

    exitProgramm = True
    break

time.sleep(2)
print(new_str)
print("Все, на этом моя работа окончена. ")
if critikal_Error != True:
    time.sleep(delayer)
    print("Надеюсь, подготовленная мной информация поможет вам в работе!")
time.sleep(delayer)
print("Счастливо!")
print(new_str)
print(new_str)
print(new_str)
ext = input("Нажмите на крестик [X] в верхнем углу окна или клавишу [ENTER] для выхода из программы...")
